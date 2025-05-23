import pandas as pd
import numpy as np
import uuid
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.units import inch
import matplotlib.pyplot as plt
import seaborn as sns
import base64
import warnings
warnings.filterwarnings('ignore')

class ReportService:
    def __init__(self):
        self.reports_folder = 'reports'
        os.makedirs(self.reports_folder, exist_ok=True)
        self.generated_reports = {}
    
    def generate_report(self, train_data, test_data, model_info, training_history, params):
        """Generate analysis report"""
        try:
            report_type = params.get('report_type', 'comprehensive')
            report_format = params.get('format', 'html')  # 'html', 'pdf', 'docx'
            include_charts = params.get('include_charts', True)
            
            # Generate report ID
            report_id = str(uuid.uuid4())
            timestamp = datetime.now()
            
            # Collect report data
            report_data = self._collect_report_data(
                train_data, test_data, model_info, training_history, params
            )
            
            # Generate report content
            if report_format == 'html':
                result = self._generate_html_report(report_data, report_id, include_charts)
            elif report_format == 'pdf':
                result = self._generate_pdf_report(report_data, report_id, include_charts)
            elif report_format == 'docx':
                result = self._generate_docx_report(report_data, report_id, include_charts)
            else:
                return {'success': False, 'message': f'不支持的报表格式: {report_format}'}
            
            if result['success']:
                # Store report info
                self.generated_reports[report_id] = {
                    'id': report_id,
                    'timestamp': timestamp.isoformat(),
                    'type': report_type,
                    'format': report_format,
                    'file_path': result['file_path'],
                    'data': report_data
                }
                
                return {
                    'success': True,
                    'message': '报表生成完成',
                    'report_id': report_id,
                    'timestamp': timestamp.isoformat(),
                    'file_path': result['file_path'],
                    'download_url': f'/api/reports/download/{report_id}/{report_format}'
                }
            else:
                return result
                
        except Exception as e:
            return {'success': False, 'message': f'生成报表失败: {str(e)}'}
    
    def download_report(self, report_id, file_format):
        """Download report file"""
        try:
            if report_id not in self.generated_reports:
                return {'success': False, 'message': '报表不存在'}
            
            report_info = self.generated_reports[report_id]
            file_path = report_info['file_path']
            
            if not os.path.exists(file_path):
                return {'success': False, 'message': '报表文件不存在'}
            
            filename = f"analysis_report_{report_id[:8]}.{file_format}"
            
            return {
                'success': True,
                'file_path': file_path,
                'filename': filename
            }
            
        except Exception as e:
            return {'success': False, 'message': f'下载报表失败: {str(e)}'}
    
    def get_report_list(self):
        """Get list of generated reports"""
        reports = []
        for report_id, info in self.generated_reports.items():
            reports.append({
                'id': report_id,
                'timestamp': info['timestamp'],
                'type': info['type'],
                'format': info['format']
            })
        
        return {
            'success': True,
            'reports': sorted(reports, key=lambda x: x['timestamp'], reverse=True)
        }
    
    def _collect_report_data(self, train_data, test_data, model_info, training_history, params):
        """Collect data for report generation"""
        data = {
            'timestamp': datetime.now().isoformat(),
            'project_info': {
                'name': params.get('project_name', '机器学习预测分析项目'),
                'description': params.get('project_description', '基于机器学习的工艺参数预测分析'),
                'author': params.get('author', 'System User')
            },
            'data_info': {},
            'model_info': {},
            'performance_metrics': {},
            'training_history': training_history or [],
            'summary': {}
        }
        
        # Data information
        if train_data is not None:
            data['data_info']['train'] = {
                'shape': train_data.shape,
                'columns': train_data.columns.tolist(),
                'numeric_columns': train_data.select_dtypes(include=[np.number]).columns.tolist(),
                'missing_values': train_data.isnull().sum().to_dict(),
                'basic_stats': train_data.describe().to_dict()
            }
        
        if test_data is not None:
            data['data_info']['test'] = {
                'shape': test_data.shape,
                'columns': test_data.columns.tolist(),
                'numeric_columns': test_data.select_dtypes(include=[np.number]).columns.tolist(),
                'missing_values': test_data.isnull().sum().to_dict()
            }
        
        # Model information
        if model_info is not None:
            data['model_info'] = {
                'model_type': model_info.get('model_type', 'Unknown'),
                'model_name': model_info.get('model_name', 'Unknown'),
                'feature_columns': model_info.get('feature_columns', []),
                'target_columns': model_info.get('target_columns', []),
                'training_time': model_info.get('training_time', ''),
                'data_shape': model_info.get('data_shape', [])
            }
            
            # Add model-specific info
            if 'params' in model_info:
                data['model_info']['parameters'] = model_info['params']
            
            if 'automl_config' in model_info:
                data['model_info']['automl_config'] = model_info['automl_config']
        
        # Generate summary
        data['summary'] = self._generate_summary(data)
        
        return data
    
    def _generate_summary(self, data):
        """Generate report summary"""
        summary = {
            'total_features': 0,
            'total_targets': 0,
            'data_quality': 'Good',
            'model_performance': 'Unknown',
            'recommendations': []
        }
        
        # Data summary
        if 'train' in data['data_info']:
            train_info = data['data_info']['train']
            summary['total_features'] = len(train_info['numeric_columns'])
            summary['total_targets'] = len(data['model_info'].get('target_columns', []))
            
            # Check data quality
            missing_ratio = sum(train_info['missing_values'].values()) / (train_info['shape'][0] * train_info['shape'][1])
            if missing_ratio > 0.1:
                summary['data_quality'] = 'Poor'
                summary['recommendations'].append('数据存在较多缺失值，建议进行数据清洗')
            elif missing_ratio > 0.05:
                summary['data_quality'] = 'Fair'
                summary['recommendations'].append('数据存在少量缺失值，建议检查数据质量')
        
        # Model recommendations
        if data['model_info'].get('model_type') == 'AutoML':
            summary['recommendations'].append('使用了AutoML进行模型选择，建议验证最优模型的泛化能力')
        
        summary['recommendations'].append('建议在更多数据上验证模型性能')
        summary['recommendations'].append('考虑进行特征工程以提升模型效果')
        
        return summary
    
    def _generate_html_report(self, data, report_id, include_charts):
        """Generate HTML report"""
        try:
            html_content = self._create_html_content(data, include_charts)
            
            file_path = os.path.join(self.reports_folder, f'report_{report_id}.html')
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'file_path': file_path,
                'message': 'HTML报表生成成功'
            }
            
        except Exception as e:
            return {'success': False, 'message': f'生成HTML报表失败: {str(e)}'}
    
    def _create_html_content(self, data, include_charts):
        """Create HTML content for report"""
        html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>机器学习分析报表</title>
            <style>
                body {{
                    font-family: 'Microsoft YaHei', Arial, sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 20px;
                    background-color: #f5f5f5;
                }}
                .container {{
                    max-width: 1200px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 30px;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }}
                h1 {{
                    color: #2c3e50;
                    text-align: center;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #34495e;
                    border-left: 4px solid #3498db;
                    padding-left: 15px;
                    margin-top: 30px;
                }}
                h3 {{
                    color: #7f8c8d;
                }}
                .section {{
                    margin: 20px 0;
                    padding: 20px;
                    background-color: #f8f9fa;
                    border-radius: 5px;
                }}
                .metric {{
                    display: inline-block;
                    margin: 10px 15px;
                    padding: 10px 20px;
                    background-color: #3498db;
                    color: white;
                    border-radius: 5px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #3498db;
                    color: white;
                }}
                .summary {{
                    background-color: #e8f6f3;
                    border-left: 5px solid #1abc9c;
                    padding: 15px;
                    margin: 20px 0;
                }}
                .recommendations {{
                    background-color: #fef9e7;
                    border-left: 5px solid #f39c12;
                    padding: 15px;
                    margin: 20px 0;
                }}
                ul {{
                    padding-left: 20px;
                }}
                .timestamp {{
                    color: #7f8c8d;
                    font-size: 0.9em;
                    text-align: right;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{data['project_info']['name']}</h1>
                <p class="timestamp">生成时间: {data['timestamp']}</p>
                
                <div class="section">
                    <h2>项目概述</h2>
                    <p>{data['project_info']['description']}</p>
                    <p><strong>作者:</strong> {data['project_info']['author']}</p>
                </div>
                
                {self._generate_data_section_html(data)}
                {self._generate_model_section_html(data)}
                {self._generate_summary_section_html(data)}
                
            </div>
        </body>
        </html>
        """
        return html
    
    def _generate_data_section_html(self, data):
        """Generate data section for HTML report"""
        if not data['data_info']:
            return ""
        
        html = '<div class="section"><h2>数据分析</h2>'
        
        if 'train' in data['data_info']:
            train_info = data['data_info']['train']
            html += f"""
            <h3>训练数据</h3>
            <div class="metric">数据行数: {train_info['shape'][0]}</div>
            <div class="metric">特征数量: {train_info['shape'][1]}</div>
            <div class="metric">数值特征: {len(train_info['numeric_columns'])}</div>
            """
            
            # Missing values table
            missing_data = [(col, count) for col, count in train_info['missing_values'].items() if count > 0]
            if missing_data:
                html += """
                <h4>缺失值统计</h4>
                <table>
                    <tr><th>列名</th><th>缺失值数量</th></tr>
                """
                for col, count in missing_data[:10]:  # Show top 10
                    html += f"<tr><td>{col}</td><td>{count}</td></tr>"
                html += "</table>"
        
        if 'test' in data['data_info']:
            test_info = data['data_info']['test']
            html += f"""
            <h3>测试数据</h3>
            <div class="metric">数据行数: {test_info['shape'][0]}</div>
            <div class="metric">特征数量: {test_info['shape'][1]}</div>
            """
        
        html += '</div>'
        return html
    
    def _generate_model_section_html(self, data):
        """Generate model section for HTML report"""
        if not data['model_info']:
            return ""
        
        model_info = data['model_info']
        html = f"""
        <div class="section">
            <h2>模型信息</h2>
            <h3>基本信息</h3>
            <div class="metric">模型类型: {model_info.get('model_type', 'Unknown')}</div>
            <div class="metric">模型名称: {model_info.get('model_name', 'Unknown')}</div>
            <div class="metric">特征数量: {len(model_info.get('feature_columns', []))}</div>
            <div class="metric">目标数量: {len(model_info.get('target_columns', []))}</div>
        """
        
        if 'training_time' in model_info:
            html += f'<p><strong>训练时间:</strong> {model_info["training_time"]}</p>'
        
        # Parameters
        if 'parameters' in model_info:
            html += "<h4>模型参数</h4><ul>"
            for key, value in model_info['parameters'].items():
                html += f"<li><strong>{key}:</strong> {value}</li>"
            html += "</ul>"
        
        html += '</div>'
        return html
    
    def _generate_summary_section_html(self, data):
        """Generate summary section for HTML report"""
        summary = data['summary']
        
        html = f"""
        <div class="summary">
            <h2>分析总结</h2>
            <div class="metric">总特征数: {summary['total_features']}</div>
            <div class="metric">目标变量数: {summary['total_targets']}</div>
            <div class="metric">数据质量: {summary['data_quality']}</div>
        </div>
        """
        
        if summary['recommendations']:
            html += """
            <div class="recommendations">
                <h3>建议与改进</h3>
                <ul>
            """
            for rec in summary['recommendations']:
                html += f"<li>{rec}</li>"
            html += "</ul></div>"
        
        return html
    
    def _generate_pdf_report(self, data, report_id, include_charts):
        """Generate PDF report"""
        try:
            file_path = os.path.join(self.reports_folder, f'report_{report_id}.pdf')
            
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            story.append(Paragraph(data['project_info']['name'], title_style))
            story.append(Spacer(1, 12))
            
            # Project info
            story.append(Paragraph("项目概述", styles['Heading2']))
            story.append(Paragraph(data['project_info']['description'], styles['Normal']))
            story.append(Paragraph(f"作者: {data['project_info']['author']}", styles['Normal']))
            story.append(Paragraph(f"生成时间: {data['timestamp']}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Data section
            if data['data_info']:
                story.append(Paragraph("数据分析", styles['Heading2']))
                
                if 'train' in data['data_info']:
                    train_info = data['data_info']['train']
                    story.append(Paragraph("训练数据统计", styles['Heading3']))
                    
                    data_table = [
                        ['项目', '数值'],
                        ['数据行数', str(train_info['shape'][0])],
                        ['特征数量', str(train_info['shape'][1])],
                        ['数值特征', str(len(train_info['numeric_columns']))]
                    ]
                    
                    table = Table(data_table)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 12))
            
            # Model section
            if data['model_info']:
                story.append(Paragraph("模型信息", styles['Heading2']))
                model_info = data['model_info']
                
                model_table = [
                    ['项目', '值'],
                    ['模型类型', model_info.get('model_type', 'Unknown')],
                    ['模型名称', model_info.get('model_name', 'Unknown')],
                    ['特征数量', str(len(model_info.get('feature_columns', [])))],
                    ['目标数量', str(len(model_info.get('target_columns', [])))]
                ]
                
                table = Table(model_table)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
            
            # Summary
            summary = data['summary']
            story.append(Paragraph("分析总结", styles['Heading2']))
            story.append(Paragraph(f"总特征数: {summary['total_features']}", styles['Normal']))
            story.append(Paragraph(f"目标变量数: {summary['total_targets']}", styles['Normal']))
            story.append(Paragraph(f"数据质量: {summary['data_quality']}", styles['Normal']))
            
            if summary['recommendations']:
                story.append(Paragraph("建议与改进", styles['Heading3']))
                for rec in summary['recommendations']:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            return {
                'success': True,
                'file_path': file_path,
                'message': 'PDF报表生成成功'
            }
            
        except Exception as e:
            return {'success': False, 'message': f'生成PDF报表失败: {str(e)}'}
    
    def _generate_docx_report(self, data, report_id, include_charts):
        """Generate DOCX report"""
        try:
            file_path = os.path.join(self.reports_folder, f'report_{report_id}.docx')
            
            doc = Document()
            
            # Title
            title = doc.add_heading(data['project_info']['name'], 0)
            title.alignment = 1  # Center alignment
            
            # Project info
            doc.add_heading('项目概述', level=1)
            doc.add_paragraph(data['project_info']['description'])
            doc.add_paragraph(f"作者: {data['project_info']['author']}")
            doc.add_paragraph(f"生成时间: {data['timestamp']}")
            
            # Data section
            if data['data_info']:
                doc.add_heading('数据分析', level=1)
                
                if 'train' in data['data_info']:
                    train_info = data['data_info']['train']
                    doc.add_heading('训练数据统计', level=2)
                    
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '项目'
                    hdr_cells[1].text = '数值'
                    
                    data_items = [
                        ('数据行数', str(train_info['shape'][0])),
                        ('特征数量', str(train_info['shape'][1])),
                        ('数值特征', str(len(train_info['numeric_columns'])))
                    ]
                    
                    for item, value in data_items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = item
                        row_cells[1].text = value
            
            # Model section
            if data['model_info']:
                doc.add_heading('模型信息', level=1)
                model_info = data['model_info']
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '项目'
                hdr_cells[1].text = '值'
                
                model_items = [
                    ('模型类型', model_info.get('model_type', 'Unknown')),
                    ('模型名称', model_info.get('model_name', 'Unknown')),
                    ('特征数量', str(len(model_info.get('feature_columns', [])))),
                    ('目标数量', str(len(model_info.get('target_columns', []))))
                ]
                
                for item, value in model_items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item
                    row_cells[1].text = value
            
            # Summary
            summary = data['summary']
            doc.add_heading('分析总结', level=1)
            doc.add_paragraph(f"总特征数: {summary['total_features']}")
            doc.add_paragraph(f"目标变量数: {summary['total_targets']}")
            doc.add_paragraph(f"数据质量: {summary['data_quality']}")
            
            if summary['recommendations']:
                doc.add_heading('建议与改进', level=2)
                for rec in summary['recommendations']:
                    p = doc.add_paragraph(rec)
                    p.style = 'List Bullet'
            
            # Save document
            doc.save(file_path)
            
            return {
                'success': True,
                'file_path': file_path,
                'message': 'DOCX报表生成成功'
            }
            
        except Exception as e:
            return {'success': False, 'message': f'生成DOCX报表失败: {str(e)}'} 